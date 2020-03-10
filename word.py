#读取docx中的文本代码示例
import docx
import re


def getContent(path):
	text=""
	file=docx.Document(path)#获取文档对象
	print("段落数："+str(len(file.paragraphs)))#段落数为几，每个回车算作一段
	#输出每一段的内容
	for para in file.paragraphs:
		text+=para.text
	return text.replace('\n','')


	
if __name__=="__main__":
	print(getContent('课程考试试卷1.docx'))


